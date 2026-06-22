"""
data_context.py — reusable source-ingestion + derivation engine (MAF: data-context)
════════════════════════════════════════════════════════════════════════════════════
Factored out of form-fill (stages B–F) so any agent/skill can build ONE unified,
queryable picture of the data a user supplies — natural-language prose AND uploaded
files (Excel/CSV/PDF/Word/images) — and then DERIVE the value a question actually
asks for (count / sum / avg / min / max with filters & group-bys, cross-tabs,
ratios, date math) rather than hoping it is written verbatim.

Everything carries PROVENANCE (file → sheet → column/row). Category vocabularies
are normalized as they are ingested (the form's words rarely match the data's).

Public surface:
    ctx = DataContext.from_sources([...]); ctx.add_text("...")
    ctx.tables / ctx.text_facts                      # the unified picture
    ctx.resolve_dimension("Female")                  # form category → (column, values)
    ctx.crosstab(rows="gender", cols="department")   # the central derive op
    ctx.derive(op="count", filters={...})            # single aggregate w/ audit
    reconcile_matrix(grid)                            # totals tie out, or flag
"""

from __future__ import annotations

import os
import re
import json
from dataclasses import dataclass, field

import pandas as pd


# ── OCR engine (live Tesseract for scanned PDFs / images) ────────────────────
def _resolve_tesseract():
    """Locate the Tesseract OCR engine binary across the TESSERACT_CMD env override,
    PATH, and the common Windows/Unix install locations. When found, points
    pytesseract at it so callers don't need it on PATH. Returns the path or None."""
    import shutil
    candidates = [
        os.environ.get("TESSERACT_CMD"),
        shutil.which("tesseract"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract", "/usr/local/bin/tesseract", "/opt/homebrew/bin/tesseract",
    ]
    for c in candidates:
        if c and os.path.exists(c):
            try:
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = c
            except Exception:
                pass
            return c
    return None


def ocr_pdf(path, dpi=300, max_pages=50):
    """Live OCR for scanned/flat PDFs with no extractable text layer. Rasterizes each
    page with pypdfium2 at `dpi` and runs Tesseract OCR. Returns (text, meta) where
    meta records the engine, pages OCR'd and per-page char counts. Raises RuntimeError
    if the OCR engine/libraries are unavailable so the caller can flag (never guess)."""
    tcmd = _resolve_tesseract()
    if not tcmd:
        raise RuntimeError("Tesseract engine not found (set TESSERACT_CMD or add it to PATH)")
    import pytesseract
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(path)
    scale = dpi / 72.0
    chunks, per_page = [], []
    n = min(len(pdf), max_pages)
    for i in range(n):
        pil = pdf[i].render(scale=scale).to_pil()
        txt = pytesseract.image_to_string(pil) or ""
        chunks.append(txt)
        per_page.append(len(txt.strip()))
    text = "\n".join(chunks).strip()
    return text, {"engine": f"tesseract@{tcmd}", "pages_ocred": n,
                  "per_page_chars": per_page, "dpi": dpi}


# ── vocabulary normalization ─────────────────────────────────────────────────
GENDER_SYNONYMS = {
    "female": {"female", "f", "w", "woman", "women", "fem", "girl", "lady"},
    "male": {"male", "m", "man", "men", "boy", "gentleman"},
}
# generic code/abbreviation hints for common department-style buckets
_DEPT_HINTS = {
    "production": {"production", "prod", "prd", "manufacturing", "mfg", "line", "factory"},
    "warehouse": {"warehouse", "wh", "ware", "whse", "logistics", "storage", "distribution", "dc"},
    "office": {"office", "off", "admin", "administration", "hq", "corporate", "back office"},
    "sales": {"sales", "sale", "commercial", "retail"},
    "quality": {"quality", "qa", "qc", "quality assurance"},
}


def _norm(s) -> str:
    return re.sub(r"\s+", " ", str(s).strip().lower()) if s is not None else ""


def match_category(form_category: str, source_value, extra_synonyms=None) -> bool:
    """Does a source cell value belong to the form's category? Handles synonyms,
    codes, casing and trailing spaces. Conservative — unknowns return False."""
    fc = _norm(form_category)
    sv = _norm(source_value)
    if not sv:
        return False
    if fc == sv:
        return True
    # explicit synonym tables
    for table in (GENDER_SYNONYMS, _DEPT_HINTS, extra_synonyms or {}):
        syns = table.get(fc)
        if syns and sv in syns:
            return True
    # prefix / code heuristics (e.g. "PROD" ~ "production", "Female"~"F")
    if len(sv) <= 4 and fc.startswith(sv):
        return True
    if sv.startswith(fc[:4]) and len(fc) >= 4:
        return True
    if fc.startswith(sv[:4]) and len(sv) >= 4:
        return True
    return False


# ── data structures ──────────────────────────────────────────────────────────
@dataclass
class Table:
    name: str               # "file.xlsx :: Sheet1"
    source: str             # file path
    sheet: str
    df: pd.DataFrame
    header_row: int = 0

    def columns(self):
        return list(self.df.columns)

    def profile(self):
        return {"name": self.name, "source": self.source, "sheet": self.sheet,
                "rows": len(self.df), "columns": list(self.df.columns),
                "dtypes": {c: str(self.df[c].dtype) for c in self.df.columns},
                "sample": self.df.head(2).to_dict("records")}


@dataclass
class DataContext:
    tables: list = field(default_factory=list)          # list[Table]
    text_facts: list = field(default_factory=list)      # list[{text, source}]
    notes: list = field(default_factory=list)

    # ---- construction -----------------------------------------------------
    @classmethod
    def from_sources(cls, sources) -> "DataContext":
        ctx = cls()
        for s in (sources or []):
            if isinstance(s, dict) and s.get("type") == "text":
                ctx.add_text(s.get("text", ""), s.get("source", "inline-text"))
            else:
                path = s["path"] if isinstance(s, dict) else s
                ctx.add_file(path)
        return ctx

    def add_text(self, text: str, source="inline-text"):
        if text and text.strip():
            self.text_facts.append({"text": text.strip(), "source": source})
        return self

    def add_file(self, path: str):
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in (".xlsx", ".xlsm", ".xls"):
                self._load_excel(path)
            elif ext == ".csv":
                self._load_csv(path)
            elif ext == ".pdf":
                self._load_pdf(path)
            elif ext in (".docx",):
                self._load_docx(path)
            elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"):
                self._load_image(path)
            else:
                self.notes.append(f"Unsupported file skipped: {path}")
        except Exception as e:
            self.notes.append(f"Failed to ingest {path}: {e!r}")
        return self

    # ---- loaders ----------------------------------------------------------
    def _load_excel(self, path):
        xl = pd.ExcelFile(path)
        for sheet in xl.sheet_names:
            raw = xl.parse(sheet, header=None)
            hr = _detect_header_row(raw)
            df = xl.parse(sheet, header=hr)
            df = df.dropna(axis=1, how="all").dropna(axis=0, how="all")
            df.columns = [str(c).strip() for c in df.columns]
            self.tables.append(Table(f"{os.path.basename(path)} :: {sheet}", path, sheet, df, hr))

    def _load_csv(self, path):
        raw = pd.read_csv(path, header=None, dtype=str, keep_default_na=False)
        hr = _detect_header_row(raw)
        df = pd.read_csv(path, header=hr)
        df.columns = [str(c).strip() for c in df.columns]
        self.tables.append(Table(f"{os.path.basename(path)} :: csv", path, "csv", df, hr))

    def _load_pdf(self, path):
        import pdfplumber
        text_chunks, ntables = [], 0
        with pdfplumber.open(path) as pdf:
            for pi, page in enumerate(pdf.pages):
                t = page.extract_text() or ""
                if t.strip():
                    text_chunks.append(t)
                for ti, tbl in enumerate(page.extract_tables() or []):
                    if tbl and len(tbl) > 1:
                        df = pd.DataFrame(tbl[1:], columns=[str(c).strip() if c else f"col{j}"
                                                            for j, c in enumerate(tbl[0])])
                        self.tables.append(Table(f"{os.path.basename(path)} :: p{pi+1}t{ti+1}",
                                                 path, f"page{pi+1}", df, 0))
                        ntables += 1
        if text_chunks:
            self.add_text("\n".join(text_chunks), source=os.path.basename(path))
        if not text_chunks and ntables == 0:
            # Scanned/flat PDF: no text layer. Run live OCR rather than just flagging it.
            try:
                ocr_text, meta = ocr_pdf(path)
                if ocr_text:
                    self.add_text(ocr_text, source=f"{os.path.basename(path)} (OCR)")
                    self.notes.append(f"{path}: OCR extracted {len(ocr_text)} chars from "
                                      f"{meta['pages_ocred']} page(s) via {meta['engine']}.")
                else:
                    self.notes.append(f"{path}: OCR produced no text (blank/illegible scan).")
            except Exception as e:
                self.notes.append(f"{path}: no text layer and OCR unavailable ({e}) "
                                  f"— flagged, not guessed.")

    def _load_docx(self, path):
        import docx
        d = docx.Document(path)
        text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        if text.strip():
            self.add_text(text, source=os.path.basename(path))
        for ti, tbl in enumerate(d.tables):
            rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
            if rows and len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=[c or f"col{j}" for j, c in enumerate(rows[0])])
                self.tables.append(Table(f"{os.path.basename(path)} :: table{ti+1}",
                                         path, f"table{ti+1}", df, 0))

    def _load_image(self, path):
        try:
            if not _resolve_tesseract():
                raise RuntimeError("Tesseract engine not found")
            import pytesseract
            from PIL import Image
            txt = pytesseract.image_to_string(Image.open(path))
            if txt.strip():
                self.add_text(txt, source=f"{os.path.basename(path)} (OCR)")
            else:
                self.notes.append(f"{path}: OCR produced no text.")
        except Exception as e:
            self.notes.append(f"{path}: image OCR unavailable ({e}) — flagged, not guessed.")

    # ---- profiling --------------------------------------------------------
    def profile(self):
        return {"tables": [t.profile() for t in self.tables],
                "text_facts": self.text_facts, "notes": self.notes}

    def primary_table(self) -> "Table":
        """The largest table — the usual roster/dataset to derive from."""
        if not self.tables:
            return None
        return max(self.tables, key=lambda t: len(t.df) * max(len(t.df.columns), 1))

    # ---- dimension resolution (form category ↔ source column) -------------
    def resolve_dimension(self, form_categories, table: "Table" = None):
        """Find the source column whose values best match the given form categories.
        Returns {column, value_map:{category:[source_values]}, table, ambiguous, missing}."""
        table = table or self.primary_table()
        if table is None:
            return {"column": None, "missing": True, "reason": "no tabular data"}
        cats = [form_categories] if isinstance(form_categories, str) else list(form_categories)
        best = None
        scores = []
        for col in table.df.columns:
            vals = [v for v in table.df[col].dropna().unique()]
            covered = 0
            vmap = {c: [] for c in cats}
            for c in cats:
                hits = [v for v in vals if match_category(c, v)]
                if hits:
                    covered += 1
                    vmap[c] = hits
            scores.append((covered, col, vmap))
        scores.sort(reverse=True, key=lambda x: x[0])
        if not scores or scores[0][0] == 0:
            return {"column": None, "missing": True,
                    "reason": f"no column matches categories {cats}", "table": table.name}
        top = scores[0]
        ambiguous = len(scores) > 1 and scores[1][0] == top[0] and scores[1][0] > 0
        return {"column": top[1], "value_map": top[2], "table": table.name,
                "ambiguous": ambiguous,
                "ambiguous_with": (scores[1][1] if ambiguous else None),
                "missing": False}

    # ---- the central derive op: cross-tab ---------------------------------
    def crosstab(self, row_categories, col_categories, table: "Table" = None,
                 row_label="rows", col_label="cols"):
        """Compute a full row×col count cross-tab in ONE pass, mapping each source
        row into the form's vocabulary. Returns the grid, totals, the expression for
        each cell, contributing row indices, and unknown/excluded counts."""
        table = table or self.primary_table()
        rres = self.resolve_dimension(row_categories, table)
        cres = self.resolve_dimension(col_categories, table)
        problems = []
        if rres.get("missing"):
            problems.append(f"row dimension unresolved: {rres.get('reason')}")
        if cres.get("missing"):
            problems.append(f"col dimension unresolved: {cres.get('reason')}")
        if rres.get("ambiguous"):
            problems.append(f"row dimension ambiguous ({rres['column']} vs {rres['ambiguous_with']})")
        if cres.get("ambiguous"):
            problems.append(f"col dimension ambiguous ({cres['column']} vs {cres['ambiguous_with']})")
        if rres.get("missing") or cres.get("missing"):
            return {"ok": False, "problems": problems, "row_res": rres, "col_res": cres}

        df = table.df
        rcol, ccol = rres["column"], cres["column"]
        rcats = [row_categories] if isinstance(row_categories, str) else list(row_categories)
        ccats = [col_categories] if isinstance(col_categories, str) else list(col_categories)

        grid = {rc: {cc: 0 for cc in ccats} for rc in rcats}
        cells_rows = {rc: {cc: [] for cc in ccats} for rc in rcats}
        unknown = 0
        for idx, rec in df.iterrows():
            rc = next((c for c in rcats if match_category(c, rec[rcol])), None)
            cc = next((c for c in ccats if match_category(c, rec[ccol])), None)
            if rc is None or cc is None:
                unknown += 1
                continue
            grid[rc][cc] += 1
            cells_rows[rc][cc].append(int(idx))
        row_tot = {rc: sum(grid[rc].values()) for rc in rcats}
        col_tot = {cc: sum(grid[rc][cc] for rc in rcats) for cc in ccats}
        grand = sum(row_tot.values())
        exprs = {rc: {cc: f"COUNT({rcol} in {{{rc}}} AND {ccol} in {{{cc}}}) = {grid[rc][cc]}"
                      for cc in ccats} for rc in rcats}
        return {"ok": True, "grid": grid, "row_totals": row_tot, "col_totals": col_tot,
                "grand_total": grand, "exprs": exprs, "cell_rows": cells_rows,
                "row_col": rcol, "col_col": ccol, "table": table.name,
                "unknown_excluded": unknown, "record_count": len(df), "problems": problems,
                "value_map": {"rows": rres["value_map"], "cols": cres["value_map"]}}

    # ---- single aggregate -------------------------------------------------
    def derive(self, op="count", column=None, filters=None, table: "Table" = None):
        """op ∈ count/sum/avg/min/max with category filters {col_or_category: value}.
        Returns the value, the expression, and the contributing row indices."""
        table = table or self.primary_table()
        if table is None:
            return {"ok": False, "reason": "no tabular data"}
        df = table.df
        mask = pd.Series([True] * len(df), index=df.index)
        applied = []
        for key, want in (filters or {}).items():
            col = key if key in df.columns else None
            if col is None:  # treat key as a category to resolve
                res = self.resolve_dimension(want if isinstance(want, str) else key, table)
                col = res.get("column")
            if col is None:
                return {"ok": False, "reason": f"filter dimension '{key}' not found", "missing": True}
            m = df[col].apply(lambda v: match_category(want, v))
            mask &= m
            applied.append(f"{col} in {{{want}}}")
        sub = df[mask]
        if op == "count":
            val = int(len(sub))
        else:
            if column is None or column not in df.columns:
                return {"ok": False, "reason": f"measure column '{column}' not found"}
            nums = pd.to_numeric(sub[column], errors="coerce").dropna()
            val = {"sum": nums.sum, "avg": nums.mean, "min": nums.min, "max": nums.max}[op]()
            val = float(val) if len(nums) else 0.0
        expr = f"{op.upper()}({'*' if op=='count' else column}" + \
               (f" WHERE {' AND '.join(applied)}" if applied else "") + f") = {val}"
        return {"ok": True, "value": val, "expr": expr, "rows": [int(i) for i in sub.index],
                "table": table.name}


# ── header-row detection (messy sheets with titles/blank rows above) ──────────
def _detect_header_row(raw: pd.DataFrame, scan=12) -> int:
    best_row, best_score = 0, -1
    n = min(scan, len(raw))
    for i in range(n):
        row = raw.iloc[i]
        nonnull = row.notna().sum()
        strings = sum(1 for v in row if isinstance(v, str) and str(v).strip())
        distinct = len(set(str(v).strip().lower() for v in row if pd.notna(v)))
        # a header row: many non-null string cells, mostly distinct, and the row
        # BELOW it has data (more numeric / filled)
        below_fill = raw.iloc[i + 1].notna().sum() if i + 1 < len(raw) else 0
        score = strings + distinct * 0.5 + (1 if below_fill >= nonnull * 0.6 else 0) * 2 - (nonnull - strings)
        if nonnull >= 2 and score > best_score:
            best_score, best_row = score, i
    return best_row


# ── reconciliation (mandatory before any fill) ────────────────────────────────
def reconcile_matrix(crosstab_result: dict) -> dict:
    """Internal-consistency checks for a derived matrix: cells sum to row totals,
    column totals and the grand total; grand total == record_count − unknown."""
    if not crosstab_result.get("ok"):
        return {"reconciled": False, "reason": "crosstab failed", "checks": []}
    grid = crosstab_result["grid"]
    row_tot = crosstab_result["row_totals"]
    col_tot = crosstab_result["col_totals"]
    grand = crosstab_result["grand_total"]
    checks, ok = [], True

    for rc, cols in grid.items():
        s = sum(cols.values())
        passed = s == row_tot[rc]
        ok &= passed
        checks.append({"check": f"row '{rc}' cells sum to row total", "expected": row_tot[rc],
                       "got": s, "pass": passed})
    for cc in col_tot:
        s = sum(grid[rc][cc] for rc in grid)
        passed = s == col_tot[cc]
        ok &= passed
        checks.append({"check": f"col '{cc}' cells sum to col total", "expected": col_tot[cc],
                       "got": s, "pass": passed})
    sum_rows = sum(row_tot.values()); sum_cols = sum(col_tot.values())
    for label, val in (("row totals", sum_rows), ("col totals", sum_cols)):
        passed = val == grand
        ok &= passed
        checks.append({"check": f"{label} sum to grand total", "expected": grand,
                       "got": val, "pass": passed})
    # grand total accounts for every record (minus reported exclusions)
    rec = crosstab_result.get("record_count")
    unk = crosstab_result.get("unknown_excluded", 0)
    if rec is not None:
        passed = grand == rec - unk
        ok &= passed
        checks.append({"check": "grand total == records − excluded/unknown",
                       "expected": rec - unk, "got": grand, "pass": passed,
                       "records": rec, "excluded_unknown": unk})
    return {"reconciled": ok, "checks": checks,
            "grand_total": grand, "excluded_unknown": unk}
