"""
form_fill.py — fill PDF / Word forms from any data source, deriving values as needed
═════════════════════════════════════════════════════════════════════════════════════
Meta App Factory | Native Python Ecosystem

Three hard parts, all handled:
  1. UNDERSTAND THE FORM — detect what is fillable and what each field means across
     PDF AcroForms, flat/scanned PDFs (coordinate overlay), and Word .docx
     (content controls / {{placeholders}} / "Label: ____" / empty table cells),
     including TABLE/MATRIX field groups captured as one structured target.
  2. UNDERSTAND THE SOURCES — via the reusable `data_context` module (prose + Excel/
     CSV/PDF/Word/images, with provenance and vocabulary normalization).
  3. DERIVE THE VALUE THE FIELD ASKS FOR — direct copy OR a computation (count/sum/
     avg/cross-tab) over the full dataset; reconcile before writing; never fabricate.

Fill preserves the original document (AcroForm values + appearances; docx cells /
placeholders; flat-PDF text overlay). A fill REPORT accompanies every output: each
filled field → value, each derived value → its formula + source + reconciliation,
everything left blank, everything left for a human, and excluded/unknown counts.

QA (non-negotiable): recompute + read back the saved file, re-check reconciliation,
render pages to images, confirm each value sits in the right row×column.
"""

from __future__ import annotations

import os
import re
import sys
import json

import pypdf

# Reuse the data-context module (stages B–F: ingest + derive + reconcile). It lives
# in its own sibling skill dir so any agent/skill can reuse it without form-fill.
_DC_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data-context")
if _DC_DIR not in sys.path:
    sys.path.insert(0, _DC_DIR)
from data_context import DataContext, reconcile_matrix, match_category  # noqa: E402

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
DATE_RE = re.compile(
    r"\b(\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{2,4}|"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b",
    re.IGNORECASE)
MONEY_RE = re.compile(r"\$\s?\d[\d,]*(?:\.\d{2})?")
SIGNATURE_HINTS = ("signature", "sign here", "signed", "initials", "date signed")


# ═══════════════════════════════════════════════════════════════════════════
#  A. FORM DETECTION & FIELD MODEL
# ═══════════════════════════════════════════════════════════════════════════

def detect_form(path: str) -> dict:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        reader = pypdf.PdfReader(path)
        fields = reader.get_fields()
        if fields:
            model = []
            for name, f in fields.items():
                model.append({"name": name, "type": str(f.get("/FT")),
                              "value": f.get("/V"),
                              "options": f.get("/Opt"),
                              "is_signature": any(h in name.lower() for h in SIGNATURE_HINTS)})
            return {"kind": "pdf_acroform", "path": path, "fields": model, "n_fields": len(model)}
        return {"kind": "pdf_flat", "path": path,
                "note": "no form fields — fill by coordinate overlay against text/table anchors"}
    if ext == ".docx":
        import docx
        d = docx.Document(path)
        text = "\n".join(p.text for p in d.paragraphs)
        placeholders = re.findall(r"\{\{[^}]+\}\}", text)
        labeled_blanks = re.findall(r"([A-Za-z][\w /]+?):\s*_{2,}", text)
        return {"kind": "docx", "path": path, "n_tables": len(d.tables),
                "placeholders": placeholders, "labeled_blanks": labeled_blanks}
    return {"kind": "unknown", "path": path}


# ═══════════════════════════════════════════════════════════════════════════
#  Natural-language fact extraction (direct values, with source snippet)
# ═══════════════════════════════════════════════════════════════════════════

def extract_nl_facts(text: str) -> dict:
    facts = {}

    def add(key, val, snippet):
        if val and key not in facts:
            facts[key] = {"value": val.strip(), "source_snippet": snippet.strip()[:120]}

    if not text:
        return facts
    for m in EMAIL_RE.finditer(text):
        add("email", m.group(0), text[max(0, m.start() - 30):m.end() + 10]); break
    for m in DATE_RE.finditer(text):
        add("date", m.group(0), text[max(0, m.start() - 25):m.end() + 10]); break
    for m in MONEY_RE.finditer(text):
        add("amount", m.group(0), text[max(0, m.start() - 30):m.end() + 10]); break
    # cue-based fields
    cues = {
        "name": r"(?:onboard|hire|employee(?:'s)? name|new hire|candidate|name)\s*[:\-]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})",
        "department": r"(?:in the|join(?:ing)?|department|dept|team|division)\s*[:\-]?\s*(?:the\s+)?([A-Z][A-Za-z]+)\b",
        "manager": r"(?:report(?:s|ing)? to|manager|reports to|supervisor)\s*[:\-]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})",
        "title": r"(?:as (?:a|an)|role|title|position)\s*[:\-]?\s*([A-Z][A-Za-z ]{2,30})",
    }
    for key, pat in cues.items():
        m = re.search(pat, text)
        if m:
            add(key, m.group(1), text[max(0, m.start() - 5):m.end() + 5])
    return facts


# ═══════════════════════════════════════════════════════════════════════════
#  B. FILLERS
# ═══════════════════════════════════════════════════════════════════════════

def fill_pdf_acroform(path: str, values: dict, out_path: str) -> str:
    reader = pypdf.PdfReader(path)
    writer = pypdf.PdfWriter()
    writer.append(reader)
    str_values = {k: ("" if v is None else str(v)) for k, v in values.items()}
    for page in writer.pages:
        try:
            writer.update_page_form_field_values(page, str_values, auto_regenerate=False)
        except Exception:
            pass
    # ensure viewers regenerate appearances so values are visible
    try:
        writer.set_need_appearances_writer(True)
    except Exception:
        pass
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    with open(out_path, "wb") as f:
        writer.write(f)
    return os.path.abspath(out_path)


def read_pdf_fields(path: str) -> dict:
    reader = pypdf.PdfReader(path)
    fields = reader.get_fields() or {}
    return {name: (f.get("/V") if hasattr(f, "get") else None) for name, f in fields.items()}


def detect_grid_anchors_pdf(path: str, row_cats, col_cats, page_index=0) -> dict:
    """Locate a matrix on a flat PDF from the table's ACTUAL cell rectangles (via
    pdfplumber.find_tables) — not word-center guesses. Maps the header row to column
    categories and the first column to row categories, then returns the exact centre
    of each interior cell so values land dead-centre in the right row×column."""
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        page = pdf.pages[page_index]
        H = page.height
        tables = page.find_tables() or []
        if not tables:
            return {"page_height": H, "col_x": {}, "row_y": {}, "total_x": None, "total_y": None}
        # pick the table whose header row best covers the requested col categories
        def cover(tbl):
            txt = tbl.extract()
            if not txt:
                return -1
            header = [(_c or "").strip() for _c in txt[0]]
            return sum(1 for cc in col_cats if any(match_category(cc, h) for h in header))
        tbl = max(tables, key=cover)
        text = tbl.extract()
        rows_geom = tbl.rows
        header = [(c or "").strip() for c in text[0]]
        first_col = [((r[0] or "").strip() if r else "") for r in text]

        def cell_center(i, j):
            bb = rows_geom[i].cells[j]
            if bb is None:
                return None
            x0, top, x1, bottom = bb
            return ((x0 + x1) / 2, (top + bottom) / 2)

        col_j = {cc: next((j for j, h in enumerate(header) if match_category(cc, h)), None) for cc in col_cats}
        row_i = {rc: next((i for i, h in enumerate(first_col) if match_category(rc, h)), None) for rc in row_cats}
        tot_j = next((j for j, h in enumerate(header) if h.lower() == "total"), None)
        tot_i = next((i for i, h in enumerate(first_col) if h.lower() == "total"), None)

        col_x, row_y = {}, {}
        # take an interior reference cell to get a column's x-centre / a row's y-centre
        ref_row = next((i for i in row_i.values() if i is not None), 1)
        ref_col = next((j for j in col_j.values() if j is not None), 1)
        for cc, j in col_j.items():
            if j is not None:
                cen = cell_center(ref_row, j)
                if cen:
                    col_x[cc] = cen[0]
        for rc, i in row_i.items():
            if i is not None:
                cen = cell_center(i, ref_col)
                if cen:
                    row_y[rc] = cen[1]
        tot_x = cell_center(ref_row, tot_j)[0] if tot_j is not None else None
        tot_y = cell_center(tot_i, ref_col)[1] if tot_i is not None else None
        return {"page_height": H, "col_x": col_x, "row_y": row_y,
                "total_x": tot_x, "total_y": tot_y}


def overlay_text_pdf(path: str, placements: list, out_path: str, font="Helvetica", size=10) -> str:
    """placements: [{page,x,y_top,text,center}] where y_top is distance from page top."""
    from reportlab.pdfgen import canvas
    from io import BytesIO
    reader = pypdf.PdfReader(path)
    writer = pypdf.PdfWriter()
    # group placements by page
    by_page = {}
    for p in placements:
        by_page.setdefault(p.get("page", 0), []).append(p)
    for i, page in enumerate(reader.pages):
        if i in by_page:
            buf = BytesIO()
            w = float(page.mediabox.width); h = float(page.mediabox.height)
            c = canvas.Canvas(buf, pagesize=(w, h))
            c.setFont(font, size)
            for pl in by_page[i]:
                y = h - pl["y_top"]
                if pl.get("center", True):
                    c.drawCentredString(pl["x"], y - size * 0.35, str(pl["text"]))
                else:
                    c.drawString(pl["x"], y - size * 0.35, str(pl["text"]))
            c.save(); buf.seek(0)
            overlay = pypdf.PdfReader(buf).pages[0]
            page.merge_page(overlay)
        writer.add_page(page)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    with open(out_path, "wb") as f:
        writer.write(f)
    return os.path.abspath(out_path)


def fill_docx(path: str, placeholders: dict, out_path: str) -> str:
    """Replace {{key}} and 'Label: ____' patterns in paragraphs and table cells,
    preserving surrounding text/styling."""
    import docx
    d = docx.Document(path)

    def patch_text(t: str) -> str:
        for k, v in placeholders.items():
            t = t.replace("{{" + k + "}}", str(v))
            t = re.sub(r"(\b" + re.escape(k) + r"\s*:\s*)_{2,}", r"\g<1>" + str(v), t, flags=re.IGNORECASE)
        return t

    def patch_para(p):
        if "{{" in p.text or re.search(r":\s*_{2,}", p.text):
            new = patch_text(p.text)
            if new != p.text:
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = new
                else:
                    p.add_run(new)

    for p in d.paragraphs:
        patch_para(p)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    patch_para(p)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    d.save(out_path)
    return os.path.abspath(out_path)


def fill_docx_matrix(path: str, grid: dict, row_cats, col_cats, out_path: str,
                     row_totals=None, col_totals=None, grand_total=None) -> dict:
    """Find the matrix table (header row matches col_cats, first column matches
    row_cats) and write each computed value into the correct row×col cell."""
    import docx
    d = docx.Document(path)
    target = None
    for tbl in d.tables:
        rows = tbl.rows
        if len(rows) < 2:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        col_idx = {cc: next((j for j, h in enumerate(header) if match_category(cc, h)), None) for cc in col_cats}
        first_col = [r.cells[0].text.strip() for r in rows]
        row_idx = {rc: next((i for i, h in enumerate(first_col) if match_category(rc, h)), None) for rc in row_cats}
        if all(v is not None for v in col_idx.values()) and all(v is not None for v in row_idx.values()):
            target = (tbl, col_idx, row_idx, header, first_col)
            break
    if target is None:
        return {"ok": False, "reason": "no matrix table matching the requested rows/cols"}
    tbl, col_idx, row_idx, header, first_col = target
    written = []
    for rc in row_cats:
        for cc in col_cats:
            cell = tbl.rows[row_idx[rc]].cells[col_idx[cc]]
            cell.text = str(grid[rc][cc])
            written.append({"row": rc, "col": cc, "value": grid[rc][cc],
                            "cell": f"r{row_idx[rc]}c{col_idx[cc]}"})
    # totals if the table has a Total row/column
    tot_col = next((j for j, h in enumerate(header) if h.strip().lower() == "total"), None)
    tot_row = next((i for i, h in enumerate(first_col) if h.strip().lower() == "total"), None)
    if tot_col is not None and row_totals:
        for rc in row_cats:
            tbl.rows[row_idx[rc]].cells[tot_col].text = str(row_totals[rc])
    if tot_row is not None and col_totals:
        for cc in col_cats:
            tbl.rows[tot_row].cells[col_idx[cc]].text = str(col_totals[cc])
    if tot_row is not None and tot_col is not None and grand_total is not None:
        tbl.rows[tot_row].cells[tot_col].text = str(grand_total)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    d.save(out_path)
    return {"ok": True, "written": written, "out_path": os.path.abspath(out_path)}


def read_docx_table_grid(path: str) -> list:
    import docx
    d = docx.Document(path)
    return [[[c.text.strip() for c in r.cells] for r in t.rows] for t in d.tables]


# ═══════════════════════════════════════════════════════════════════════════
#  B2. GOOGLE DOCS — fill a live Google Doc directly via the Docs API
# ═══════════════════════════════════════════════════════════════════════════
# Same fill model as the .docx path ({{key}} placeholders), but instead of
# exporting a local file we authenticate to Google, fetch the document, and push
# edits straight into the live doc with a single batchUpdate. The request set is
# built by a pure function (`build_replace_requests`) so the logic is verifiable
# offline (dry_run) without credentials or network.

GOOGLE_DOCS_SCOPES = ["https://www.googleapis.com/auth/documents"]


def _google_docs_service(credentials_path: str = None, token_path: str = None):
    """OAuth (installed-app) flow for the Google Docs API. Reuses a cached token,
    refreshes it silently when expired, and only opens the consent browser flow on
    first use. Paths default to the GOOGLE_OAUTH_CLIENT_SECRET / GOOGLE_OAUTH_TOKEN
    env vars, else credentials.json / token.json in the cwd."""
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build

    credentials_path = credentials_path or os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET", "credentials.json")
    token_path = token_path or os.environ.get("GOOGLE_OAUTH_TOKEN", "token.json")

    creds = None
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, GOOGLE_DOCS_SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(credentials_path):
                raise FileNotFoundError(
                    f"Google OAuth client secret not found at '{credentials_path}'. "
                    "Download it from the Google Cloud console (OAuth 2.0 Client ID, "
                    "Desktop app) and point GOOGLE_OAUTH_CLIENT_SECRET at it.")
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, GOOGLE_DOCS_SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, "w", encoding="utf-8") as f:
            f.write(creds.to_json())
    return build("docs", "v1", credentials=creds)


def read_google_doc_text(document: dict) -> str:
    """Flatten a Docs API document resource to plain text (paragraph text runs and
    table cell text), so we can read back/verify what the live doc contains."""
    def _walk(content):
        out = []
        for el in content or []:
            para = el.get("paragraph")
            if para:
                for pe in para.get("elements", []):
                    tr = pe.get("textRun")
                    if tr:
                        out.append(tr.get("content", ""))
            table = el.get("table")
            if table:
                for row in table.get("tableRows", []):
                    for cell in row.get("tableCells", []):
                        out.append(_walk(cell.get("content", [])))
        return "".join(out)
    return _walk(document.get("body", {}).get("content", []))


def build_replace_requests(placeholders: dict) -> list:
    """Pure: build the Docs API `replaceAllText` request list for {{key}} placeholders.
    Separated so the request construction is unit-testable without auth/network."""
    reqs = []
    for k, v in placeholders.items():
        reqs.append({"replaceAllText": {
            "containsText": {"text": "{{" + str(k) + "}}", "matchCase": True},
            "replaceText": "" if v is None else str(v)}})
    return reqs


def fill_google_doc(doc_id: str, placeholders: dict, credentials_path: str = None,
                    token_path: str = None, dry_run: bool = False) -> dict:
    """Fill a live Google Doc by replacing {{key}} placeholders with values, in place.
    Authenticates, fetches the document content, pushes one batchUpdate, then reads
    the doc back to confirm the placeholders are gone and the values are present.
    dry_run=True returns the request set without touching the network (offline proof)."""
    requests = build_replace_requests(placeholders)
    if dry_run:
        return {"status": "dry_run", "doc_id": doc_id, "n_requests": len(requests),
                "requests": requests}

    service = _google_docs_service(credentials_path, token_path)
    before = service.documents().get(documentId=doc_id).execute()
    before_text = read_google_doc_text(before)

    resp = service.documents().batchUpdate(
        documentId=doc_id, body={"requests": requests}).execute()
    replies = resp.get("replies", [])
    changed = {}
    for k, r in zip(placeholders.keys(), replies):
        changed[k] = r.get("replaceAllText", {}).get("occurrencesChanged", 0)

    after = service.documents().get(documentId=doc_id).execute()
    after_text = read_google_doc_text(after)

    # QA read-back: every value present, no {{placeholders}} left for the filled keys
    values_present = {k: (str(v) in after_text) for k, v in placeholders.items() if v not in (None, "")}
    leftover = re.findall(r"\{\{[^}]+\}\}", after_text)
    report = {
        "status": "filled",
        "doc_id": doc_id,
        "doc_url": f"https://docs.google.com/document/d/{doc_id}/edit",
        "title": after.get("title"),
        "occurrences_changed": changed,
        "filled": [{"placeholder": k, "value": v, "occurrences": changed.get(k, 0)}
                   for k, v in placeholders.items()],
        "qa": {"clean": all(values_present.values()) and not leftover,
               "values_present": values_present,
               "remaining_placeholders": leftover,
               "chars_before": len(before_text), "chars_after": len(after_text)},
    }
    return report


# ═══════════════════════════════════════════════════════════════════════════
#  C. FILL REPORT
# ═══════════════════════════════════════════════════════════════════════════

class FillReport:
    def __init__(self, form: dict):
        self.form = {"kind": form.get("kind"), "path": form.get("path")}
        self.filled = []
        self.derived = []
        self.left_blank = []
        self.human_action_required = []
        self.excluded_unknown = 0
        self.reconciliation = None
        self.out_path = None
        self.qa = None

    def add_direct(self, field, value, source):
        self.filled.append({"field": field, "value": value, "kind": "direct", "source": source})

    def add_derived(self, field, value, expression, source_columns, source_table, reconciliation=None):
        self.filled.append({"field": field, "value": value, "kind": "derived",
                            "derivation": expression, "source": {"table": source_table,
                            "columns": source_columns}})
        self.derived.append({"target": field, "value": value, "expression": expression,
                            "source_columns": source_columns, "source_table": source_table,
                            "reconciliation": reconciliation})

    def blank(self, field, reason):
        self.left_blank.append({"field": field, "reason": reason})

    def human(self, field, reason):
        self.human_action_required.append({"field": field, "reason": reason})

    def to_dict(self):
        return {"form": self.form, "filled": self.filled, "derived": self.derived,
                "left_blank": self.left_blank, "human_action_required": self.human_action_required,
                "excluded_unknown_rows": self.excluded_unknown, "reconciliation": self.reconciliation,
                "out_path": self.out_path, "qa": self.qa}

    def text(self):
        lines = [f"FILL REPORT — {self.form['kind']} — {os.path.basename(self.form['path'] or '')}",
                 f"  Output: {self.out_path}"]
        lines.append(f"  Filled fields ({len(self.filled)}):")
        for f in self.filled:
            if f["kind"] == "derived":
                lines.append(f"    - {f['field']} = {f['value']}   [DERIVED]")
                lines.append(f"        {f['derivation']}")
                lines.append(f"        source: {f['source']['table']} cols {f['source']['columns']}")
            else:
                lines.append(f"    - {f['field']} = {f['value']}   [direct] (from {f['source']})")
        if self.left_blank:
            lines.append("  Left blank (needs input):")
            for b in self.left_blank:
                lines.append(f"    - {b['field']} -- {b['reason']}")
        if self.human_action_required:
            lines.append("  Left for a human (on purpose):")
            for h in self.human_action_required:
                lines.append(f"    - {h['field']} -- {h['reason']}")
        if self.excluded_unknown:
            lines.append(f"  Excluded/unknown source rows: {self.excluded_unknown}")
        if self.reconciliation is not None:
            lines.append(f"  Reconciliation: {'PASS' if self.reconciliation.get('reconciled') else 'FAIL'}")
        if self.qa is not None:
            lines.append(f"  QA: {'CLEAN' if self.qa.get('clean') else 'ISSUES'} — {self.qa.get('summary')}")
        return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════
#  D. QA  (render + read-back + re-reconcile)
# ═══════════════════════════════════════════════════════════════════════════

def _render_pdf(path, resolution=120):
    """Render pages to PNG. Uses pypdfium2 WITH form drawing so AcroForm field
    values are visible in the render (pdfplumber's rasterizer ignores NeedAppearances
    and would show blank fields — always verify the *rendered* file). Falls back to
    pdfplumber if pypdfium2 is unavailable."""
    out_dir = os.path.splitext(path)[0] + "_pages"
    os.makedirs(out_dir, exist_ok=True)
    pngs = []
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(path)
        try:
            pdf.init_forms()           # enables form-field appearance drawing
        except Exception:
            pass
        scale = resolution / 72.0
        for i in range(len(pdf)):
            png = os.path.join(out_dir, f"page_{i+1}.png")
            pdf[i].render(scale=scale).to_pil().save(png)
            pngs.append(png)
        return pngs
    except Exception:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                png = os.path.join(out_dir, f"page_{i+1}.png")
                page.to_image(resolution=resolution).save(png)
                pngs.append(png)
        return pngs


# ═══════════════════════════════════════════════════════════════════════════
#  E. HIGH-LEVEL ORCHESTRATORS
# ═══════════════════════════════════════════════════════════════════════════

def fill_matrix_form(form_path: str, sources: list, row_cats, col_cats, out_path: str,
                     fill_totals=True) -> dict:
    """Central use case: derive a row×col cross-tab from uploaded data and write each
    computed value into the matrix cells of an existing form. Reconciles BEFORE
    writing; refuses to fill (and reports) if totals do not tie out."""
    form = detect_form(form_path)
    report = FillReport(form)
    ctx = DataContext.from_sources(sources)

    ct = ctx.crosstab(row_cats, col_cats)
    if not ct.get("ok"):
        report.reconciliation = {"reconciled": False, "problems": ct.get("problems")}
        for p in ct.get("problems", []):
            report.blank("matrix", p)
        return {"status": "not_filled", "report": report.to_dict(), "report_text": report.text()}

    rec = reconcile_matrix(ct)
    report.reconciliation = rec
    report.excluded_unknown = ct["unknown_excluded"]
    if not rec["reconciled"]:
        # Numbers that don't add up on a form are worse than blanks — do NOT fill.
        return {"status": "not_filled_reconciliation_failed",
                "report": report.to_dict(), "report_text": report.text(), "crosstab": ct}

    grid, exprs = ct["grid"], ct["exprs"]
    src_cols = [ct["row_col"], ct["col_col"]]
    # record each cell as a derived value (formula + source + reconciliation)
    for rc in row_cats:
        for cc in col_cats:
            report.add_derived(f"{rc} x {cc}", grid[rc][cc], exprs[rc][cc], src_cols,
                               ct["table"], reconciliation="ties to row/col/grand totals")

    if form["kind"] == "docx":
        res = fill_docx_matrix(form_path, grid, row_cats, col_cats, out_path,
                               row_totals=ct["row_totals"] if fill_totals else None,
                               col_totals=ct["col_totals"] if fill_totals else None,
                               grand_total=ct["grand_total"] if fill_totals else None)
        if not res.get("ok"):
            report.blank("matrix", res.get("reason"))
            return {"status": "not_filled", "report": report.to_dict(), "report_text": report.text()}
        report.out_path = res["out_path"]
        # QA: read the saved cells back, confirm they match the computed grid
        grids = read_docx_table_grid(res["out_path"])
        report.qa = _qa_docx_matrix(grids, grid, row_cats, col_cats, ct)

    elif form["kind"] == "pdf_flat":
        anchors = detect_grid_anchors_pdf(form_path, row_cats, col_cats)
        placements = []
        missing = [rc for rc in row_cats if rc not in anchors["row_y"]] + \
                  [cc for cc in col_cats if cc not in anchors["col_x"]]
        if missing:
            report.blank("matrix", f"could not anchor labels on flat PDF: {missing}")
            return {"status": "not_filled", "report": report.to_dict(), "report_text": report.text()}
        for rc in row_cats:
            for cc in col_cats:
                placements.append({"page": 0, "x": anchors["col_x"][cc],
                                   "y_top": anchors["row_y"][rc], "text": grid[rc][cc]})
            if fill_totals and anchors.get("total_x") is not None:
                placements.append({"page": 0, "x": anchors["total_x"],
                                   "y_top": anchors["row_y"][rc], "text": ct["row_totals"][rc]})
        if fill_totals and anchors.get("total_y") is not None:
            for cc in col_cats:
                placements.append({"page": 0, "x": anchors["col_x"][cc],
                                   "y_top": anchors["total_y"], "text": ct["col_totals"][cc]})
            if anchors.get("total_x") is not None:
                placements.append({"page": 0, "x": anchors["total_x"],
                                   "y_top": anchors["total_y"], "text": ct["grand_total"]})
        report.out_path = overlay_text_pdf(form_path, placements, out_path)
        pngs = _render_pdf(report.out_path)
        report.qa = _qa_pdf_overlay(report.out_path, grid, row_cats, col_cats, anchors, pngs)
    else:
        report.blank("matrix", f"unsupported form kind for matrix fill: {form['kind']}")
        return {"status": "not_filled", "report": report.to_dict(), "report_text": report.text()}

    return {"status": "filled", "report": report.to_dict(), "report_text": report.text(),
            "crosstab": ct}


def fill_simple_form(form_path: str, text: str, field_map: dict, out_path: str,
                     extra_facts: dict = None) -> dict:
    """Fill a form's direct fields from a natural-language paragraph.
    field_map maps form-field-name → nl-fact-key (e.g. {'EmployeeName':'name'}).
    Never fabricates: an unfound fact leaves the field blank in the report."""
    form = detect_form(form_path)
    report = FillReport(form)
    facts = extract_nl_facts(text)
    if extra_facts:
        for k, v in extra_facts.items():
            facts.setdefault(k, {"value": v, "source_snippet": "(provided)"})

    values = {}
    for field, fact_key in field_map.items():
        f = detect_is_signature(form, field)
        if f:
            report.human(field, "signature/attestation — never auto-signed")
            continue
        fact = facts.get(fact_key)
        if fact and fact["value"]:
            values[field] = fact["value"]
            report.add_direct(field, fact["value"], f"NL: \"{fact['source_snippet']}\"")
        else:
            report.blank(field, f"no '{fact_key}' found in the provided text")

    if form["kind"] == "pdf_acroform":
        report.out_path = fill_pdf_acroform(form_path, values, out_path)
        readback = read_pdf_fields(report.out_path)
        pngs = _render_pdf(report.out_path)
        ok = all(str(readback.get(k, "")) .strip() == str(v).strip() for k, v in values.items())
        report.qa = {"clean": ok, "summary": f"{len(values)} fields read back; render {len(pngs)} page(s)",
                     "png_pages": pngs, "readback": {k: str(readback.get(k)) for k in values}}
    elif form["kind"] == "docx":
        report.out_path = fill_docx(form_path, values, out_path)
        grids = read_docx_table_grid(report.out_path)
        import docx
        joined = "\n".join(p.text for p in docx.Document(report.out_path).paragraphs)
        ok = all(str(v) in joined for v in values.values())
        report.qa = {"clean": ok, "summary": f"{len(values)} placeholders written; read back from docx"}
    else:
        report.blank("(form)", f"unsupported form kind: {form['kind']}")
    return {"status": "filled" if report.out_path else "not_filled",
            "report": report.to_dict(), "report_text": report.text(), "facts": facts}


def detect_is_signature(form: dict, field_name: str) -> bool:
    if form.get("kind") == "pdf_acroform":
        for f in form.get("fields", []):
            if f["name"] == field_name:
                return f.get("is_signature", False)
    return any(h in field_name.lower() for h in SIGNATURE_HINTS)


# ── QA helpers ────────────────────────────────────────────────────────────────
def _qa_docx_matrix(grids, grid, row_cats, col_cats, ct):
    """Read the saved docx cells back and confirm each value sits in the right cell;
    re-check the matrix totals against the saved numbers."""
    issues = []
    # locate the saved matrix grid (the one whose header matches col_cats)
    saved = None
    for g in grids:
        header = g[0]
        if all(any(match_category(cc, h) for h in header) for cc in col_cats):
            saved = g; break
    if saved is None:
        return {"clean": False, "summary": "could not re-read matrix table"}
    header = saved[0]
    col_idx = {cc: next(j for j, h in enumerate(header) if match_category(cc, h)) for cc in col_cats}
    row_idx = {rc: next(i for i, r in enumerate(saved) if match_category(rc, r[0])) for rc in row_cats}
    for rc in row_cats:
        for cc in col_cats:
            got = saved[row_idx[rc]][col_idx[cc]]
            if str(got) != str(grid[rc][cc]):
                issues.append(f"{rc}×{cc}: saved '{got}' != computed {grid[rc][cc]}")
    # re-reconcile from the saved numbers
    saved_grid = {rc: {cc: int(saved[row_idx[rc]][col_idx[cc]]) for cc in col_cats} for rc in row_cats}
    row_ok = all(sum(saved_grid[rc].values()) == ct["row_totals"][rc] for rc in row_cats)
    col_ok = all(sum(saved_grid[rc][cc] for rc in row_cats) == ct["col_totals"][cc] for cc in col_cats)
    return {"clean": not issues and row_ok and col_ok,
            "summary": f"{len(row_cats)*len(col_cats)} cells read back; "
                       f"totals tie={row_ok and col_ok}",
            "issues": issues}


def _qa_pdf_overlay(path, grid, row_cats, col_cats, anchors, pngs):
    """Independently confirm each value landed in the right cell: re-detect the table
    on the FILLED pdf, re-map header/first-column labels to categories, and read each
    interior cell's text directly (NOT via the placement anchors — so a placement bug
    cannot self-validate)."""
    import pdfplumber
    issues = []
    with pdfplumber.open(path) as pdf:
        page = pdf.pages[0]
        tables = page.find_tables() or []
        grids = [t.extract() for t in tables]
    saved = None
    for g in grids:
        if g and all(any(match_category(cc, (h or "")) for h in g[0]) for cc in col_cats):
            saved = g; break
    if saved is None:
        return {"clean": False, "summary": "could not re-detect matrix on filled PDF",
                "issues": ["no table"], "png_pages": pngs}
    header = [(h or "").strip() for h in saved[0]]
    first_col = [((r[0] or "").strip() if r else "") for r in saved]
    col_j = {cc: next(j for j, h in enumerate(header) if match_category(cc, h)) for cc in col_cats}
    row_i = {rc: next(i for i, h in enumerate(first_col) if match_category(rc, h)) for rc in row_cats}
    for rc in row_cats:
        for cc in col_cats:
            got = (saved[row_i[rc]][col_j[cc]] or "").strip()
            if got != str(grid[rc][cc]):
                issues.append(f"{rc} x {cc}: cell holds '{got}', expected {grid[rc][cc]}")
    return {"clean": not issues,
            "summary": f"{len(row_cats)*len(col_cats)} cells read back from the filled PDF table; "
            f"rendered {len(pngs)} page(s)", "issues": issues, "png_pages": pngs}
