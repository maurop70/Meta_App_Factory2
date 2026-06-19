"""
samples.py — generate demo forms + roster for form-fill (fixtures / try-it)
═══════════════════════════════════════════════════════════════════════════
Produces three blank forms and a messy Excel roster so the skill can be
exercised end-to-end without external files:
  • onboarding_acroform.pdf  — PDF AcroForm with named text fields (+ signature)
  • headcount_flat.pdf       — flat PDF with a drawn matrix table (no form fields)
  • headcount_matrix.docx    — Word doc with a matrix table + {{placeholders}}
  • roster.xlsx              — employee roster with mixed vocab/codes + a title row
"""

from __future__ import annotations

import os
import pandas as pd

ROW_CATS = ["Male", "Female"]
COL_CATS = ["Production", "Warehouse", "Office"]


def make_roster(path):
    data = [
        ("A. Stone", "F", "Production"), ("B. Lee", "Female", "PROD"), ("C. Diaz", "M", "Production"),
        ("D. Ng", "Male", "PROD"), ("E. Roy", "F", "Warehouse"), ("G. Poe", "Female", "WH"),
        ("H. Ito", "male", "Warehouse"), ("I. Cruz", "M", "WH"), ("J. Kim", "F", "Office"),
        ("K. Vora", "W", "Office"), ("L. Sun", "M", "Admin"), ("M. Oja", "Female", "Production"),
        ("N. Bell", "F", "PROD"), ("O. Park", "Male", "Warehouse"), ("P. Yu", "M", "WH"),
        ("Q. Rao", "female", "Production"), ("R. Fox", "F", "Warehouse"), ("S. Ade", "Man", "Office"),
        ("T. Lim", "F", "Production"), ("U. Beck", "M", "PROD"), ("V. Cho", "Female", "Warehouse"),
        ("W. Ali", "M", "Office"), ("X. Mora", "F", "WH"), ("Y. Tan", "Male", "Production"),
        ("Z. Ray", "F", ""),   # blank dept → excluded/unknown (reported, not guessed)
    ]
    df = pd.DataFrame(data, columns=["Employee Name", "Sex", "Dept"])
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        pd.DataFrame([["ACME Manufacturing — Q2 Headcount Roster (CONFIDENTIAL)"]]).to_excel(
            w, sheet_name="Roster", index=False, header=False, startrow=0)
        df.to_excel(w, sheet_name="Roster", index=False, startrow=2)
    return os.path.abspath(path)


def make_onboarding_acroform(path):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import HexColor
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    c = canvas.Canvas(path, pagesize=letter)
    w, h = letter
    c.setFillColor(HexColor("#1F3864"))
    c.rect(0, h - 70, w, 70, fill=1, stroke=0)
    c.setFillColor(HexColor("#FFFFFF")); c.setFont("Helvetica-Bold", 18)
    c.drawString(54, h - 46, "New Hire Onboarding Form")
    c.setFillColor(HexColor("#000000")); c.setFont("Helvetica", 10)
    fields = [("Employee Name", "EmployeeName"), ("Start Date", "StartDate"),
              ("Department", "Department"), ("Work Email", "Email"),
              ("Annual Salary", "Salary"), ("Reporting Manager", "Manager")]
    y = h - 120
    form = c.acroForm
    for label, name in fields:
        c.setFont("Helvetica", 10)
        c.drawString(54, y + 4, label + ":")
        form.textfield(name=name, x=190, y=y - 4, width=330, height=20, borderWidth=1,
                       borderColor=HexColor("#999999"), fillColor=HexColor("#F4F7FB"),
                       fontSize=10, forceBorder=True)
        y -= 42
    c.drawString(54, y + 4, "Signature:")
    form.textfield(name="Signature", x=190, y=y - 4, width=330, height=24, borderWidth=1,
                   borderColor=HexColor("#999999"), forceBorder=True)
    c.setFont("Helvetica-Oblique", 8); c.setFillColor(HexColor("#808080"))
    c.drawString(54, 40, "Office use only — do not write below this line.")
    c.save()
    return os.path.abspath(path)


def _matrix_cells():
    header = ["Headcount"] + COL_CATS + ["Total"]
    rows = [header]
    for rc in ROW_CATS:
        rows.append([rc] + ["" for _ in COL_CATS] + [""])
    rows.append(["Total"] + ["" for _ in COL_CATS] + [""])
    return rows


def make_headcount_flat_pdf(path):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import HexColor
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    c = canvas.Canvas(path, pagesize=letter)
    w, h = letter
    c.setFillColor(HexColor("#1F3864")); c.rect(0, h - 70, w, 70, fill=1, stroke=0)
    c.setFillColor(HexColor("#FFFFFF")); c.setFont("Helvetica-Bold", 18)
    c.drawString(54, h - 46, "Workforce Headcount Report  (EEO-style)")
    c.setFillColor(HexColor("#000000")); c.setFont("Helvetica", 10)
    c.drawString(54, h - 95, "Complete the grid: number of employees by gender and department.")
    grid = _matrix_cells()
    ncol = len(grid[0]); nrow = len(grid)
    x0, top = 54, 150
    col_w = [120, 110, 110, 110, 90]
    row_h = 44
    # draw cells
    c.setFont("Helvetica", 11)
    yt = top
    for i, row in enumerate(grid):
        x = x0
        for j, cell in enumerate(row):
            cw = col_w[j]
            # header row / first col shaded
            if i == 0 or j == 0:
                c.setFillColor(HexColor("#D9E1F2")); c.rect(x, h - yt - row_h, cw, row_h, fill=1, stroke=0)
            c.setStrokeColor(HexColor("#444444")); c.setLineWidth(0.7)
            c.rect(x, h - yt - row_h, cw, row_h, fill=0, stroke=1)
            if cell:
                c.setFillColor(HexColor("#000000"))
                c.setFont("Helvetica-Bold" if (i == 0 or j == 0) else "Helvetica", 11)
                c.drawCentredString(x + cw / 2, h - yt - row_h / 2 - 4, str(cell))
            x += cw
        yt += row_h
    c.setFont("Helvetica-Oblique", 8); c.setFillColor(HexColor("#808080"))
    c.drawString(54, 60, "Prepared by HR. For official use.")
    c.save()
    return os.path.abspath(path)


def make_headcount_docx(path):
    import docx
    from docx.shared import Pt, RGBColor
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    d = docx.Document()
    d.add_heading("Workforce Headcount Report", level=0)
    d.add_paragraph("Company: {{company}}        Reporting period: {{period}}")
    d.add_paragraph("Complete the grid below: number of employees by gender and department.")
    grid = _matrix_cells()
    tbl = d.add_table(rows=len(grid), cols=len(grid[0]))
    tbl.style = "Table Grid"
    for i, row in enumerate(grid):
        for j, cell in enumerate(row):
            tbl.rows[i].cells[j].text = str(cell)
    d.add_paragraph("Prepared by: {{prepared_by}}")
    d.save(path)
    return os.path.abspath(path)


def make_all(out_dir="out"):
    return {
        "roster": make_roster(os.path.join(out_dir, "roster.xlsx")),
        "onboarding_pdf": make_onboarding_acroform(os.path.join(out_dir, "onboarding_acroform.pdf")),
        "headcount_flat_pdf": make_headcount_flat_pdf(os.path.join(out_dir, "headcount_flat.pdf")),
        "headcount_docx": make_headcount_docx(os.path.join(out_dir, "headcount_matrix.docx")),
    }


if __name__ == "__main__":
    import json
    print(json.dumps(make_all(), indent=2))
